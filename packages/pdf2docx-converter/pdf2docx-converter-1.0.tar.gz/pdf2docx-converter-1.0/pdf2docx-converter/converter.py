import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from PIL import Image
import io

def pdf_to_docx(pdf_path, docx_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document()
    
    for page_number in range(pdf_document.page_count):
        page = pdf_document.load_page(page_number)
        image_list = page.get_pixmapmatrix()
        
        text = page.get_text()
        doc.add_paragraph(text)
        
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image = Image.open(io.BytesIO(base_image["image"]))
            image.save(f"image_{page_number}_{img_index}.png")
            doc.add_picture(f"image_{page_number}_{img_index}.png", width=Inches(5))

    doc.save(docx_path)
